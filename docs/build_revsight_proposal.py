from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_PATH = OUTPUT_DIR / "revsight-propuesta-servicios.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT = RGBColor(244, 246, 249)
BORDER = "D9E1F2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(p, before=0, after=8, line=1.33, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    p.alignment = align


def add_text_paragraph(doc, text, before=0, after=8, line=1.33, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    style_paragraph(p, before=before, after=after, line=line, align=align)
    run = p.add_run(text)
    set_run_font(run, size=11, color=TEXT)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        style_paragraph(p, before=18, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
        size = 16
        color = BLUE
    else:
        style_paragraph(p, before=12, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
        size = 13
        color = DARK
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, before=0, after=4, line=1.2, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(item)
        set_run_font(run, size=11, color=TEXT)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("Revsight | Propuesta de Servicios")
    set_run_font(run, size=9, color=MUTED, italic=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(fp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    fr = fp.add_run("Documento de propuesta comercial")
    set_run_font(fr, size=9, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    style_paragraph(p, before=28, after=8, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("REVSIGHT")
    set_run_font(run, size=14, color=MUTED, bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("Propuesta de Prestacion de Servicios")
    set_run_font(run, size=24, color=TEXT, bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=22, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("Revenue Management, Distribucion Hotelera y Soporte Funcional de Sistemas")
    set_run_font(run, size=12, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.15), Inches(4.1)]
    rows = [
        ("Propiedades", "Wyndham Garden Los Mochis | Esplendor by Wyndham Los Mochis"),
        ("Modalidad", "Esquema externo remoto / hibrido"),
        ("Honorario propuesto", "MXN $22,000.00 + IVA mensuales"),
        ("Vigencia inicial", "90 dias con revision de alcance y carga operativa"),
    ]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if j == 0:
                set_cell_shading(cell, "F4F6F9")
            p = cell.paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(rows[i][j])
            set_run_font(run, size=11, color=TEXT, bold=(j == 0))

    doc.add_paragraph()
    p = doc.add_paragraph()
    style_paragraph(p, before=16, after=6, line=1.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("Presenta")
    set_run_font(run, size=10.5, color=MUTED, italic=True)
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("Revsight")
    set_run_font(run, size=13, color=TEXT, bold=True)


def add_scope_table(doc):
    add_heading(doc, "Alcance General del Servicio", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.15), Inches(4.1)]
    headers = ["Modulo", "Alcance principal"]
    for j, cell in enumerate(table.rows[0].cells):
        cell.width = widths[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "EAF1FB")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        style_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(headers[j])
        set_run_font(run, size=10.5, color=TEXT, bold=True)

    data = [
        ("Revenue Management", "Analisis de demanda, ocupacion, ADR, RevPAR, pacing, pick-up y recomendaciones tarifarias."),
        ("Distribucion Hotelera", "Seguimiento de OTAs, paridad, inventario, disponibilidad y configuraciones comerciales."),
        ("Sistemas", "Soporte funcional sobre Opera Cloud PMS y SynXis CRS, incluyendo revision de incidencias y configuraciones criticas."),
        ("Seguimiento Comercial", "Reportes ejecutivos, hallazgos, riesgos y priorizacion de pendientes dentro del alcance acordado."),
    ]
    for left, right in data:
        row = table.add_row()
        for j, text in enumerate((left, right)):
            cell = row.cells[j]
            cell.width = widths[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(text)
            set_run_font(run, size=10.5, color=TEXT, bold=(j == 0))


def add_pricing_table(doc):
    add_heading(doc, "Propuesta Economica", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(2.4), Inches(3.85)]
    rows = [
        ("Honorario mensual", "MXN $22,000.00 + IVA"),
        ("Cobertura", "Ambas propiedades"),
        ("Esquema", "Prestacion de servicios externos principalmente remotos / hibridos"),
        ("Revision", "A los 90 dias para validar carga operativa, continuidad y ajustes"),
    ]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if j == 0:
                set_cell_shading(cell, "F4F6F9")
            p = cell.paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(rows[i][j])
            set_run_font(run, size=10.5, color=TEXT, bold=(j == 0))


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)
add_header_footer(section)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
styles["Normal"].font.size = Pt(11)

add_cover(doc)
doc.add_page_break()

add_heading(doc, "Resumen Ejecutivo", level=1)
add_text_paragraph(
    doc,
    "La presente propuesta tiene como objetivo brindar continuidad tecnica y comercial a Wyndham Garden Los Mochis y Esplendor by Wyndham Los Mochis en las areas de revenue management, distribucion hotelera, soporte funcional de sistemas y seguimiento de procesos criticos vinculados a inventario, tarifas, canales y comercializacion.",
)
add_text_paragraph(
    doc,
    "El valor de este esquema radica en sostener funciones sensibles de la operacion bajo un modelo externo con alcance definido, seguimiento estructurado y claridad de responsabilidades, evitando dependencia improvisada sobre tareas que impactan ocupacion, ADR, RevPAR, configuraciones, canales y continuidad comercial.",
)
add_text_paragraph(
    doc,
    "El servicio podra apoyarse en herramientas propias de analisis y seguimiento desarrolladas dentro de Revsight, con el fin de aportar mayor estructura, trazabilidad y consistencia al proceso de revision, reporte y recomendacion.",
)

add_heading(doc, "Objetivo del Servicio", level=1)
add_bullets(
    doc,
    [
        "Brindar soporte especializado en revenue management, estrategia tarifaria y lectura de demanda.",
        "Dar seguimiento funcional a distribucion hotelera, OTAs, inventario, paridad y configuraciones comerciales.",
        "Acompanhar incidencias y validaciones sobre Opera Cloud PMS y SynXis CRS dentro del alcance acordado.",
        "Emitir reportes y recomendaciones ejecutivas para apoyar la toma de decisiones comerciales y operativas.",
    ],
)

add_scope_table(doc)

add_heading(doc, "Modelo de Trabajo", level=1)
add_bullets(
    doc,
    [
        "Esquema principal remoto con intervenciones hibridas puntuales cuando se requiera y se acuerde previamente.",
        "Revsight emitira reportes, hallazgos y recomendaciones ejecutivas periodicas.",
        "Las recomendaciones tarifarias seran canalizadas a la persona interna responsable de aplicar cambios en tarifas.",
        "Las incidencias se atenderan con base en prioridad, alcance y ventana razonable de servicio.",
    ],
)

doc.add_page_break()

add_heading(doc, "Entregables", level=1)
add_bullets(
    doc,
    [
        "Seguimiento recurrente de revenue y distribucion para ambas propiedades.",
        "Recomendaciones ejecutivas sobre tarifas, inventario, canales y riesgos comerciales.",
        "Reportes de Revsight con analisis, observaciones y acciones sugeridas para seguimiento interno.",
        "Uso de herramientas propias de analisis y seguimiento de Revsight como apoyo metodologico para reportes y recomendaciones.",
        "Soporte funcional sobre Opera Cloud PMS y SynXis CRS dentro del alcance acordado.",
    ],
)

add_heading(doc, "Exclusiones del Servicio", level=1)
add_bullets(
    doc,
    [
        "No sustituye funciones de Gerencia General, operacion diaria ni administracion de personal.",
        "No incluye disponibilidad permanente ni cobertura 24/7.",
        "No incluye ejecucion operativa ilimitada fuera del alcance definido.",
        "No incluye la implementacion interna continua de cambios tarifarios, salvo acuerdo especifico por separado.",
    ],
)

add_pricing_table(doc)

add_heading(doc, "Condiciones y Siguientes Pasos", level=1)
add_bullets(
    doc,
    [
        "Definir una contraparte interna autorizada para seguimiento y validacion de decisiones.",
        "Iniciar con una etapa de 90 dias para estabilizacion, continuidad y revision de carga operativa real.",
        "Revisar al cierre del periodo la continuidad del esquema, posibles ajustes de alcance y condiciones economicas.",
        "Formalizar el servicio bajo esquema externo de prestacion de servicios profesionales.",
        "El uso de herramientas propias de Revsight forma parte del metodo de trabajo del prestador y no implica cesion de software, licencia independiente, desarrollo adicional ni transferencia de propiedad intelectual, salvo acuerdo expreso por separado.",
    ],
)

add_text_paragraph(
    doc,
    "Revsight presenta esta propuesta con el objetivo de aportar continuidad, estructura y acompanamiento especializado en areas clave para la estabilidad comercial y tecnica de ambas propiedades, con un modelo claro de responsabilidades y alcance.",
    before=12,
    after=0,
)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
